import os, re, json, unicodedata, uuid, random, time
from collections import defaultdict, deque
from typing import Dict, Any, List, Set, Optional

from fastapi import FastAPI, Request, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from dotenv import load_dotenv

try:
    from openai import OpenAI
    _OPENAI_IMPORTED = True
except Exception:
    OpenAI = None 
    _OPENAI_IMPORTED = False

from docx import Document

DOCX_PRODUCTOS: List[str] = []

try:
    doc = Document("data/catalogo.docx")
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            DOCX_PRODUCTOS.append(text)
    print(f"[INFO] Catálogo DOCX cargado con {len(DOCX_PRODUCTOS)} entradas")
except Exception as e:
    print(f"[WARN] No se pudo cargar catalogo.docx: {e}")



# ---------------- SETUP ----------------
load_dotenv()

# Config
APP_ENV = os.getenv("APP_ENV", "dev")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
STATIC_DIR = os.getenv("STATIC_DIR", "public")
PRODUCTS_PATH = os.getenv("PRODUCTS_PATH", "productos.json")
DOCX_PATH = os.getenv("DOCX_PATH", "data/catalogo.docx")
CORS_ALLOWED = [o.strip() for o in os.getenv("CORS_ALLOWED_ORIGINS", "*").split(",") if o.strip()]
RATE_LIMIT_PER_MIN = int(os.getenv("RATE_LIMIT_PER_MIN", "60"))
MAX_MSG_LEN = int(os.getenv("MAX_MSG_LEN", "800"))

client: Optional[OpenAI] = None
if _OPENAI_IMPORTED and OPENAI_API_KEY:
    try:
        client = OpenAI(api_key=OPENAI_API_KEY)
    except Exception:
        client = None

conversation_state: Dict[str, dict] = {}
rate_buckets: Dict[str, deque] = defaultdict(deque)

# ---------------- APP ----------------
app = FastAPI(title="Chatbot Ecolite API", version="1.3.1")
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ALLOWED,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

if os.path.isdir(STATIC_DIR):
    app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")


# ---------------- UTILIDADES ----------------
def _norm(s: Any) -> str:
    if s is None:
        return ""
    s = str(s).lower()
    return "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")


def _ip(request: Request) -> str:
    fwd = request.headers.get("x-forwarded-for")
    if fwd:
        return fwd.split(",")[0].strip()
    client_host = request.client.host if request.client else "unknown"
    return client_host or "unknown"


def _rate_limit(key: str):
    now = time.time()
    bucket = rate_buckets[key]
    while bucket and now - bucket[0] > 60:
        bucket.popleft()
    if len(bucket) >= RATE_LIMIT_PER_MIN:
        raise HTTPException(status_code=429, detail="Rate limit exceeded")
    bucket.append(now)


def log_event(event: str, **kwargs):
    try:
        payload = {"event": event, **kwargs, "ts": round(time.time(), 3)}
        print(json.dumps(payload, ensure_ascii=False))
    except Exception:
        pass


# ---------------- DATA -------------------
PRODUCTOS: Dict[str, dict] = {}


def _load_productos(path: str) -> Dict[str, dict]:
    if not os.path.exists(path):
        log_event("products_missing", path=path)
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            if not isinstance(data, dict):
                return {}
            cleaned = {}
            for code, p in data.items():
                if not isinstance(p, dict):
                    continue
                cleaned[code] = {
                    "name": p.get("name"),
                    "price": p.get("price"),
                    "img_url": p.get("img_url"),
                    "url": p.get("url"),
                    "tags": p.get("tags", []),
                    "categories": p.get("categories", []),
                }
            return cleaned
    except Exception as e:
        log_event("products_load_error", error=str(e))
        return {}


PRODUCTOS = _load_productos(PRODUCTS_PATH)


# -------- Embeddings --------
from openai import OpenAI
import numpy as np

EMBED_MODEL = "text-embedding-3-small"
client = OpenAI(api_key=OPENAI_API_KEY)

PRODUCT_EMBEDS: dict[str, np.ndarray] = {}

def build_embeddings():
    texts = []
    codes = []
    for code, p in PRODUCTOS.items():
        desc = f"{p.get('name','')} {','.join(p.get('tags',[]))} {','.join(p.get('categories',[]))}"
        texts.append(desc)
        codes.append(code)

    batch_size = 100
    for i in range(0, len(texts), batch_size):
        chunk = texts[i:i+batch_size]
        resp = client.embeddings.create(model=EMBED_MODEL, input=chunk)
        for j, emb in enumerate(resp.data):
            PRODUCT_EMBEDS[codes[i+j]] = np.array(emb.embedding)

    print(f"[INFO] Embeddings construidos: {len(PRODUCT_EMBEDS)} productos")

if PRODUCTOS and client:
    build_embeddings()

def buscar_por_embeddings(query: str, k: int = 5) -> list[dict]:
    """Busca productos similares usando embeddings de OpenAI"""
    if not PRODUCT_EMBEDS:
        return []

    q_emb = client.embeddings.create(model=EMBED_MODEL, input=[query]).data[0].embedding
    q_vec = np.array(q_emb)

    sims = []
    for code, emb in PRODUCT_EMBEDS.items():
        sim = np.dot(q_vec, emb) / (np.linalg.norm(q_vec) * np.linalg.norm(emb))
        sims.append((sim, code))

    sims.sort(reverse=True)  
    

    out = []
    for _, code in sims[:k]:
        p = PRODUCTOS[code]
        out.append({
            "code": code,
            "name": p.get("name"),
            "price": p.get("price"),
            "img_url": p.get("img_url"),
            "url": p.get("url"),
            "tags": p.get("tags", []),
            "categories": p.get("categories", []),
        })
    return out


# --- DOCX: carga robusta + texto unificado ---
DOCX_TEXT = ""
try:
    if os.path.exists(DOCX_PATH):
        _doc = Document(DOCX_PATH)  # usar 'Document', no 'DocxDocument'
        DOCX_PRODUCTOS = []
        for p in _doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                DOCX_PRODUCTOS.append(t)
        DOCX_TEXT = "\n".join(DOCX_PRODUCTOS)
    else:
        DOCX_PRODUCTOS = DOCX_PRODUCTOS  # ya poblado arriba si cargó por ruta fija
except Exception:
    DOCX_PRODUCTOS = DOCX_PRODUCTOS
    DOCX_TEXT = "\n".join(DOCX_PRODUCTOS)



# -------------- PROMPTS ----------------
SYSTEM_PROMPT_CONDUCTOR = """
Eres asesor de iluminación de Ecolite. Responde en 1–2 frases, tono WhatsApp, claro y amable (acento colombiano).
Guía SIEMPRE este flujo, sin decir que existe un flujo:

FORMULARIO (datos)
1) ESPACIO (ambiente/lugar requiere la iluminación)
   - Si falta, pregúntalo: “¿En qué espacio sería la instalación (ej: oficina, pasillo, piscina…)?”

2) ¿QUIERES QUE TE SUGIERA O TIENES UNA LUMINARIA EN ESPECÍFICO?
   - Si ya hay ESPACIO y no hay MODO, pregunta: “¿Quieres que te sugiera o tienes una luminaria en específico?”
   - Si el usuario dice que QUIERE SUGERENCIAS → action=SHOW_SUGGESTIONS (5 productos aleatorios del ESPACIO).
   - Si el usuario dice que TIENE UNA EN ESPECÍFICO o da código/nombre → action=SEARCH_SPECIFIC.

3) SI → QUIERE QUE LE SUGIERA (después de mostrar 5)
   - Pregunta: “¿Te sirve alguna de las luminarias sugeridas o estás buscando algo diferente?”
   - Si contesta que SÍ → action=COLLECT_SPECS (pide, en ese orden, Temperatura, Vatios y Tipo de instalación; una sola pregunta por turno).
   - Si dice que busca diferente → action=ASK_MODE.

4) SI → TIENE UNA EN ESPECÍFICO
   - action=SEARCH_SPECIFIC.

Reglas:
- NUNCA dispares SHOW_SUGGESTIONS hasta que el usuario responda que QUIERE SUGERENCIAS.
- Si el usuario pide “más”, usa action=MORE_SUGGESTIONS y el reply exacto: "Te muestro 5 más 👌"
- Evita repetir el espacio (“oficina”, “piscina”) si ya quedó claro.
- No agregues preguntas extra en el mismo turno de MORE_SUGGESTIONS.
- NUNCA inventes productos. El backend se encarga de mostrar resultados.
- NO uses markdown ni listas.

Devuelve SOLO JSON con:
{
  "reply": string,
  "action": "ASK_SPACE" | "ASK_MODE" | "SHOW_SUGGESTIONS" | "ASK_AFTER_SUGGESTIONS" | "COLLECT_SPECS" | "SEARCH_SPECIFIC" | "MORE_SUGGESTIONS" | "SMALLTALK" | "NONE",
  "space": string|null,
  "mode": "sugerir"|"especifico"|null,
  "spec_field": "temperatura"|"vatios"|"instalacion"|null
}
"""


# -------------- NORM & GUARDS ----------
code_regex   = re.compile(r"\b(?=[A-Z0-9-]{3,15}\b)(?=[A-Z0-9-]*\d)[A-Z0-9-]+\b")
CODE_SOFT_RE = re.compile(r"[A-Z]{2,}\d+[A-Z0-9-]*", re.I)

# Heurísticas para intención
SUGERIR_RE = re.compile(r"\b(sug|sugi|sugie|sugier|sugiere|sugi[eé]reme|sugerir|sugerencia|recomiend|recomienda|sugiere|recomiende|sugiera|recomi[eé]ndame)\w*\b", re.I)
ESPECIFICO_RE = re.compile(r"\b(especific|espec[ií]fica|tengo\s+una|ya\s+tengo|c[oó]digo|codigo|referencia|sku|modelo|nombre)\w*\b", re.I)
def said_suggest(txt: str) -> bool: return bool(SUGERIR_RE.search(txt or ""))
def said_specific(txt: str) -> bool: return bool(ESPECIFICO_RE.search(txt or ""))

CHEAP_RE  = re.compile(
    r"\b(m[aá]s\s+barat[oa]s?|barat[oa]s?|econ[oó]mic[oa]s?|menor(\s+de)?\s+precio|m[aá]s\s+econ[oó]mic[oa]s?|baratic[ao]s?|low\s*cost)\b",
    re.I,
)
CATEGORY_SYNONYMS: dict[str, set[str]] = {
    "panel": [
        "panel", "paneles", "panel led", "paneles led", "panel 60x60", "600x600",
        "panel cuadrado", "panel redondo", "panel backlight", "panel slim",
        "panel sobreponer", "panel incrustar"
    ],
        "tubo": [
        "tubo", "tubos", "tubo led", "tubos led", "tubos leds", "fluorescente", "fluorescentes",
        "t8", "t5", "tube light", "led tube"
    ],
    "driver": [
        "driver", "fuente", "fuentes", "fuente de poder", "fuentes de poder",
        "source", "transformador", "adaptador", "alimentador", "power supply"
    ],
    "lamp_decorativa": [
        "lampara decorativa","Lámpara Decorativa de Techo", "lamparas decorativas",
        "lámpara decorativa", "lámparas decorativas",
        "lampara de techo decorativa", "lamparas de techo decorativas",
        "lámpara decorativa de techo", "lámparas decorativas de techo"
],
    "lamp": [
        "lámpara", "lampara", "lámparas", "lamparas",
        "bombillo", "bombillos", "foco", "focos", "ampolleta", "ampolletas", "bulbo", "bulbos", "spot", "spots"
    ],
    "seguridad": [
        "cámara", "camara", "cámaras", "camaras", "seguridad", "cctv",
        "videocámara", "videocamara", "videovigilancia", "camara ip",
        "camara wifi", "camara solar", "domo", "dome", "bullet"
    ],
    "exterior": [
        "reflector", "reflectores", "proyector", "proyectores",
        "floodlight", "spot exterior", "exterior", "jardín", "jardin",
        "fachada", "parqueadero", "calle", "vía", "vias", "vial"
    ],    
    "alumbrado_publico": [
        "alumbrado publico", "alumbrado público", "poste", "postes",
        "farola", "farolas", "luz publica", "luz pública",
        "calle", "carretera", "vía", "vias", "vial", "avenida"
    ],
    "interior": [
        "downlight", "downlights", "plafón", "plafones",
        "empotrable", "sobremuro", "colgante", "techo", "oficina",
        "pasillo", "sala", "dormitorio", "interior"
    ],
    "emergencia": [
        "emergencia", "emergencias", "luz emergencia", "luces emergencia",
        "aviso salida", "exit", "salida", "emergency light", "backup"
    ],
    "solar": [
        "solar", "panel solar", "reflector solar", "lampara solar",
        "foco solar", "proyector solar", "kit solar", "cargador solar"
    ],
    "decorativa": [
        "decorativa", "decorativas", "cinta led", "tira led", "strip led",
        "rgb", "neón", "neon", "flexible", "manguera luminosa"
    ],
    "industrial": [
        "highbay", "high bay", "campana", "campanas",
        "nave", "bodega", "planta", "industrial", "naves industriales"
    ],
    "accesorio": [
        "accesorio", "accesorios", "soporte", "soportes",
        "riel", "rieles", "magnético", "magnetico", "track",
        "perfil", "perfiles", "difusor", "clips", "conector", "conectores"
    ]
}
def want_cheapest(txt: str) -> bool: return bool(CHEAP_RE.search(txt or ""))

def detect_category(txt: str):
    t = _norm(txt)
    for key in sorted(CATEGORY_SYNONYMS.keys(), key=lambda k: -len(k)):
        for s in CATEGORY_SYNONYMS[key]:
            if _norm(s) in t:
                return key
    return None




CHEAPEST_Q_RE = re.compile(
    r"\b(cu[aá]l(\s+es)?\s+el\s+m[aá]s\s+barat[oa]s?)\b|\b(qu[eé]\s+es\s+lo\s+m[aá]s\s+barat[oa])\b",
    re.I,
)
def is_cheapest_question(txt: str) -> bool:
    """
    True si el usuario pregunta explícitamente por el más barato.
    También cubre variantes sin tilde ('cual', 'que').
    """
    if not txt: return False
    if CHEAPEST_Q_RE.search(txt): return True
    t_norm = _norm(txt)
    return bool(want_cheapest(txt) and re.search(r"\b(cual|que)\b", t_norm))


# -------------- RAG LIGERO (señales) -------------
SIGNALS = {
    "piscina": {"piscina", "sumergible", "ip68", "bajo agua", "bajo_agua"},
    "oficina": {"oficina", "panel", "lineal", "downlight"},
    "exterior": {"exterior", "ip65", "ip66", "fachada", "terraza", "jardin", "jardín", "proyector", "aplique", "poste"},
}
DOCX_TEXT_N = _norm(DOCX_TEXT)
if DOCX_TEXT_N:
    if "piscina" in DOCX_TEXT_N: SIGNALS["piscina"] |= {"nicho", "pentair", "12v", "rgb"}
    if "oficina" in DOCX_TEXT_N: SIGNALS["oficina"] |= {"ugr", "anti deslumbramiento", "uniformidad"}
    if any(k in DOCX_TEXT_N for k in ["exterior","fachada","terraza","jardin","jardín"]): SIGNALS["exterior"] |= {"baliza"}

def es_pregunta_de_cantidad(txt: str) -> bool:
    """
    Detecta si el usuario está preguntando por más opciones
    sin mencionar un código concreto.
    """
    t = _norm(txt)
    expresiones = [
        "solo tienes", "único", "única", "hay más", "otra opción", 
        "no hay más", "más modelos", "más opciones"
    ]
    return any(e in t for e in expresiones)




def must_from_text(txt: str) -> Set[str]:
    t = _norm(txt)
    must: Set[str] = set()
    if any(s in t for s in SIGNALS["piscina"]):  must |= {"piscina","sumergible"}
    if any(s in t for s in SIGNALS["oficina"]):  must |= {"oficina"}
    if any(s in t for s in SIGNALS["exterior"]): must |= {"ip65"}
    return must


def extraer_codigo(txt: str) -> Optional[str]:
    """
    Devuelve un posible código de producto en el mensaje.
    Acepta tokens alfanuméricos (mínimo 4 caracteres),
    pero ignora palabras comunes como 'MUESTRAME', 'SUGIERE', etc.
    """
    if not txt:
        return None

    q = txt.upper()

    matches = re.findall(r"\b[A-Z0-9-]{4,15}\b", q)
    if not matches:
        return None

    stopwords = {"MUESTRAME", "MUESTRA", "ENSÉÑAME", "SUGIERE", "SUGIEREME"}

    valid = [m for m in matches if m not in stopwords]

    if not valid:
        return None

    for m in valid:
        if m in PRODUCTOS:
            return m

    valid.sort(key=lambda m: len(m), reverse=True)
    return valid[0]




def buscar_producto_especifico(query_text: str, space: Optional[str] = None, k: int = 5) -> List[dict]:
    q = (query_text or "").strip()
    qn = _norm(q)
    q_upper = q.upper()

    code_token = extraer_codigo(q_upper)
    code_prefix = None
    if code_token:
        m_pref = re.match(r"([A-Z]+)", code_token)
        code_prefix = m_pref.group(1) if m_pref else None

    name_tokens = [w for w in re.split(r"\s+", qn) if len(w) > 2]
    piscina_mode = (space and _norm(space) == "piscina")

    def score_item(code: str, prod: dict) -> int:
        name = prod.get("name") or ""
        name_n = _norm(name)
        tags_n = _norm(" ".join(map(str, prod.get("tags", []))))
        cats_n = _norm(" ".join(map(str, prod.get("categories", []))))
        full_n = f"{name_n} {tags_n} {cats_n}"

        if piscina_mode and not (("sumergible" in full_n) or ("ip68" in full_n)):
            return 0

        s = 0
        cu = code.upper()

        if code_token:
            if cu == code_token: s += 100
            elif cu.startswith(code_token): s += 70
            elif code_token in cu: s += 45

        if code_prefix:
            if cu.startswith(code_prefix): s += 35
            else: s -= 10

        for t in name_tokens:
            if t in name_n: s += 14
            if t in tags_n: s += 7
            if t in cats_n: s += 4

        if space:
            sp = _norm(space)
            if sp in full_n: s += 8
            if sp == "piscina" and (("sumergible" in full_n) or ("ip68" in full_n)): s += 10

        return s

    candidatos: List[tuple[int, str]] = []
    for code, p in PRODUCTOS.items():
        sc = score_item(code, p)
        if sc > 0: candidatos.append((sc, code))

    candidatos.sort(reverse=True)
    out = []
    for _, code in candidatos[:k]:
        p = PRODUCTOS[code]
        out.append({
            "code": code, "name": p.get("name"), "price": p.get("price"),
            "img_url": p.get("img_url"), "url": p.get("url"),
            "tags": p.get("tags", []), "categories": p.get("categories", []),
        })
    return out


# -------------- Precio helpers -------------
def _digits_to_int(s: str) -> Optional[int]:
    digits = "".join(ch for ch in s if ch.isdigit())
    return int(digits) if digits else None

def _price_value(price) -> Optional[int]:
    if isinstance(price, (int, float)): return int(price)
    if isinstance(price, str):
        parts = re.split(r"[-–—]", price)
        vals = []
        for p in parts:
            v = _digits_to_int(p)
            if v is not None: vals.append(v)
        if vals: return min(vals)
    return None

def _prod_category_tokens(p) -> set[str]:
    toks = set()
    for cat in p.get("categories", []) or []:
        for tk in _norm(cat).split():
            tt = _normalize_token(tk)
            if tt and tt not in {"de","del","para","y"}:
                toks.add(tt)
    return toks

def _cat_to_tokens(cat_key: str) -> set[str]:
    if cat_key == "lamp_decorativa":
        return {"lampara","decorativa"}
    if cat_key == "lamp":
        return {"lampara"}
    if cat_key in CATEGORY_SYNONYMS:
        toks = set()
        for phr in CATEGORY_SYNONYMS[cat_key]:
            for tk in _norm(phr).split():
                toks.add(_normalize_token(tk))
        return {t for t in toks if t not in {"de","del","para","y"}}
    return {_normalize_token(t) for t in _norm(cat_key).split() if t}

def product_is_in_category(p, cat_key_or_text: str) -> bool:
    """Devuelve True si el producto pertenece a la categoría pedida,
    comprobándolo contra name, tags o categories."""
    if not cat_key_or_text:
        return True

    req_tokens = _cat_to_tokens(cat_key_or_text)

    texto = _norm(
        (p.get("name") or "") + " " +
        " ".join(p.get("tags", [])) + " " +
        " ".join(p.get("categories", []))
    )

    return all(tok in texto for tok in req_tokens)

# -------------- BUSCADOR (general) ---------------
def buscar_productos(
    query_text: str,
    space: Optional[str] = None,
    k: int = 5,
    must_any: Optional[Set[str]] = None,
    must_have: Optional[Set[str]] = None,
    exclude_codes: Optional[Set[str]] = None,
    cheapest_first: bool = False,
) -> List[Dict[str, Any]]:
    """
    Busca productos en PRODUCTOS aplicando filtros opcionales.
    - must_any: al menos uno de estos tokens debe aparecer
    - must_have: todos estos tokens deben aparecer (en name, tags o categories)
    - exclude_codes: códigos a ignorar
    - cheapest_first: ordena por precio ascendente
    """

    resultados: List[Dict[str, Any]] = []
    exclude_codes = exclude_codes or set()

    for code, p in PRODUCTOS.items():
        if code in exclude_codes:
            continue

        norm_name = _norm(p.get("name", ""))
        norm_tags = [_norm(t) for t in p.get("tags", [])]
        norm_cats = [_norm(c) for c in p.get("categories", [])]

        if must_have:
            if not all(
                (m in norm_name) or (m in norm_tags) or (m in norm_cats)
                for m in must_have
            ):
                continue

        if must_any:
            if not any(
                (m in norm_name) or (m in norm_tags) or (m in norm_cats)
                for m in must_any
            ):
                continue

        resultados.append({
            "code": code,
            "name": p.get("name"),
            "price": p.get("price"),
            "img_url": p.get("img_url"),
            "url": p.get("url"),
            "tags": p.get("tags", []),
            "categories": p.get("categories", []),
        })

    # --- Ordenar ---
    if cheapest_first:
        resultados.sort(key=lambda x: x.get("price") or 1e9)
    else:
        resultados.sort(key=lambda x: x.get("name") or "")

    return resultados[:k]



# -------------- Parsers de SPECs / presupuesto ----------
def parse_temperatura(s: str) -> Optional[str]:
    t = _norm(s)
    if any(w in t for w in ["calida","cálida","warm","3000","2700","3000k","2700k"]): return "cálida"
    if any(w in t for w in ["neutra","neutral","4000","4000k"]): return "neutra"
    if any(w in t for w in ["fria","fría","cool","6000","6500","6000k","6500k"]): return "fría"
    return None

def parse_vatios(s: str) -> Optional[int]:
    m = re.search(r"(\d{1,3})\s*w", _norm(s))
    return int(m.group(1)) if m else None

def parse_instalacion(s: str) -> Optional[str]:
    t = _norm(s)
    if "incrust" in t or "empotr" in t:
        return "incrustar"
    if "sobrep" in t or "superficie" in t:
        return "sobreponer"
    return None

_BUDGET_RE = re.compile(r"(\$|cop|col\$)?\s*([0-9][0-9\.,\s]*)(k|mil|m|millones)?", re.I)
def parse_presupuesto(s: str) -> Optional[int]:
    t = _norm(s)
    m = _BUDGET_RE.search(t)
    if not m: return None
    num = m.group(2).replace(" ", "").replace(".", "").replace(",", "")
    try: base = int(num)
    except Exception: return None
    suf = (m.group(3) or "").lower()
    if suf in {"k","mil"}: base *= 1000
    elif suf in {"m","millones"}: base *= 1_000_000
    return base


# -------------- Detección de espacio --------------
SPACE_CANON = {
    "piscina": "piscina",
    "oficina": "oficina",
    "pasillo": "pasillo",
    "bodega": "bodega",
    "sala": "sala",
    "cocina": "cocina",
    "baño": "baño",
    "banio": "baño",
    "terraza": "terraza",
    "jardin": "jardin",
    "jardín": "jardin",
    "fachada": "fachada",
    "parqueadero": "parqueadero",
    "retail": "retail",    
    "bodega": "bodega",
    "industrial": "industrial",
    "dormitorio": "dormitorio",
    "escalera": "escalera",
    "alumbrado": "alumbrado_publico",
    "alumbrado publico": "alumbrado_publico",
    "alumbrado público": "alumbrado_publico",
    "poste": "alumbrado_publico",
    "farola": "alumbrado_publico",
    "farolas": "alumbrado_publico",
    "calle": "alumbrado_publico",
    "carretera": "alumbrado_publico",
    "vía": "alumbrado_publico",
    "vias": "alumbrado_publico",
    "vial": "alumbrado_publico"
}

def _normalize_token(tok: str) -> str:
    t = _norm(tok)
    if t.endswith("es") and len(t) > 4:
        return t[:-2]
    if t.endswith("s") and len(t) > 3:
        return t[:-1]
    return t

# --- tokens dinámicos para categoría/espacio ---
def tokens_para_clave(clave: Optional[str]) -> set[str]:
    if not clave:
        return set()
    c = _norm(clave)
    toks = set()
    if c in CATEGORY_SYNONYMS:
        for phr in CATEGORY_SYNONYMS[c]:
            for tk in _norm(phr).split():
                tkn = _normalize_token(tk)
                if tkn and tkn not in {"de", "del", "para", "y", "la", "el", "los", "las"}:
                    toks.add(tkn)
    else:
        for tk in c.split():
            tkn = _normalize_token(tk)
            if tkn:
                toks.add(tkn)
    return toks

# Señales flexibles desde DOCX (se autorrellenan si el DOCX trae pistas)
DOCX_TEXT_N = _norm(DOCX_TEXT or "")
DOCX_SEÑALES = {
    "piscina": {"piscina", "sumergible", "ip68", "nicho", "rgb", "12v", "bajo", "agua"},
    "oficina": {"oficina", "panel", "ugr", "uniformidad"},
    "exterior": {"exterior", "ip65", "ip66", "baliza", "fachada", "terraza", "jardin"},
}
# Autorefuerzo: si el DOCX contiene términos, añadimos variantes
for k, base in list(DOCX_SEÑALES.items()):
    for t in list(base):
        if t in DOCX_TEXT_N:
            base.add(t)  # idempotente; sirve como semilla



# --- Índices ligeros para JSON+DOCX ---
def _producto_texto_full(p: dict) -> str:
    return _norm(
        (p.get("name") or "") + " " +
        " ".join(p.get("tags", []) or []) + " " +
        " ".join(p.get("categories", []) or [])
    )

def _score_bm25_like(texto: str, q_tokens: list[str]) -> float:
    # sencillo: conteo ponderado por log, sin dependencias extra
    if not texto or not q_tokens:
        return 0.0
    hits = sum(1 for t in q_tokens if t in texto)
    return hits + (0.2 * sum(texto.count(t) for t in q_tokens))  # denso > 1ª aparición

def _expand_with_space_and_docx(query: str, space: Optional[str]) -> str:
    extra = []
    space_toks = tokens_para_clave(space)
    if space:
        extra.extend(space_toks)
        # sumar señales del DOCX si la clave existe
        if space in DOCX_SEÑALES:
            extra.extend(list(DOCX_SEÑALES[space]))
    return (query or "").strip() + " " + " ".join(sorted(set(extra)))

def _exact_code(q: str) -> Optional[str]:
    code = extraer_codigo(q)
    return code if code and code in PRODUCTOS else None

def _hard_filters_ok(full_text: str, space: Optional[str], cat: Optional[str]) -> bool:
    # Filtros "duros" contextuales pero dinámicos:
    # 1) si hay espacio con tokens, exigir al menos 1 token
    req_space = tokens_para_clave(space)
    if req_space and not any(t in full_text for t in req_space):
        return False
    # 2) si hay categoría con tokens, exigir al menos 1
    req_cat = tokens_para_clave(cat)
    if req_cat and not any(t in full_text for t in req_cat):
        return False
    return True

def buscar_hibrido(query_text: str, space: Optional[str], cat: Optional[str], k: int = 5) -> list[dict]:
    q = (query_text or "").strip()
    qn = _norm(q)
    code = _exact_code(qn)
    if code:
        p = PRODUCTOS[code]
        return [{
            "code": code,
            "name": p.get("name"),
            "price": p.get("price"),
            "img_url": p.get("img_url"),
            "url": p.get("url"),
            "tags": p.get("tags", []),
            "categories": p.get("categories", []),
        }]
    # Expandimos consulta con espacio + señales de DOCX
    q_expanded = _expand_with_space_and_docx(qn, space)
    q_tokens = [_normalize_token(t) for t in q_expanded.split() if len(t) > 2]

    # 1) Candidatos por BM25-like del JSON
    candidatos = []
    for code, p in PRODUCTOS.items():
        full = _producto_texto_full(p)
        if not _hard_filters_ok(full, space, cat):
            continue
        s = 0.0
        # preferencia por frase exacta del nombre
        name_n = _norm(p.get("name") or "")
        if q and q in name_n:
            s += 4.0
        s += _score_bm25_like(full, q_tokens)
        candidatos.append((s, code))

    # 2) Si pocos candidatos, añadir por embeddings (top-N) y re-rank
    if len(candidatos) < k:
        try:
            # embeddings sobre consulta expandida
            emb_results = buscar_por_embeddings(q_expanded, k=min(30, max(10, k*4)))
            for pr in emb_results:
                full = _producto_texto_full(pr)
                if not _hard_filters_ok(full, space, cat):
                    continue
                # sim boost aproximado (lista ya viene ordenada por similitud)
                # usamos posición inversa para aportar puntaje decreciente
                idx = emb_results.index(pr)
                sim_boost = max(0.0, 6.0 - 0.15 * idx)
                candidatos.append((sim_boost, pr["code"]))
        except Exception:
            pass

    # 3) Re-rank consolidado (sumar duplicados, priorizar exactitud)
    agg: dict[str, float] = {}
    for s, c in candidatos:
        agg[c] = agg.get(c, 0.0) + s

    # 4) Boost por señales fuertes (piscina: ip68 / sumergible automáticas sin ifs duros)
    for c, sc in list(agg.items()):
        p = PRODUCTOS.get(c, {})
        full = _producto_texto_full(p)
        # refuerzo semántico cuando coincide la clave del espacio
        if space and any(t in full for t in tokens_para_clave(space)):
            agg[c] = sc + 1.5
        # refuerzo semántico categoría
        if cat and any(t in full for t in tokens_para_clave(cat)):
            agg[c] = agg[c] + 1.0

    ordenados = sorted(agg.items(), key=lambda x: x[1], reverse=True)
    out = []
    for code, _ in ordenados:
        p = PRODUCTOS[code]
        out.append({
            "code": code,
            "name": p.get("name"),
            "price": p.get("price"),
            "img_url": p.get("img_url"),
            "url": p.get("url"),
            "tags": p.get("tags", []),
            "categories": p.get("categories", []),
        })
        if len(out) >= k:
            break
    return out


def sugerencias_aleatorias(space=None, n=5, exclude_codes=None, categoria=None):
    exclude_codes = exclude_codes or set()
    candidatos = []

    for code, p in PRODUCTOS.items():
        if code in exclude_codes:
            continue

        if categoria:
            must_tokens = {_norm(s) for s in CATEGORY_SYNONYMS.get(categoria, [])}
            texto = _norm(
                (p.get("name") or "") + " " +
                " ".join(p.get("tags", [])) + " " +
                " ".join(p.get("categories", []))
            )

            if not any(tok in texto for tok in must_tokens):
                continue

        if space:
            texto_space = _norm(" ".join(p.get("tags", [])) + " " + (p.get("name") or ""))
            if space.lower() not in texto_space:
                continue

        candidatos.append({
            "code": code,
            "name": p.get("name"),
            "price": p.get("price"),
            "img_url": p.get("img_url"),
            "url": p.get("url"),
            "tags": p.get("tags", []),
            "categories": p.get("categories", []),
        })

    random.shuffle(candidatos)
    return [c for c in candidatos if c["code"] not in exclude_codes][:n]



def detect_space_from_text(text: str) -> Optional[str]:
    t = _norm(text)
    for k in SPACE_CANON.keys():
        if k in t: return SPACE_CANON[k]
    return None


# -------------- LLM ORCHESTRATION ----------
def _parse_json_loose(text: str) -> Optional[dict]:
    if not text: return None
    s = text.strip()
    if s.startswith("```"):
        s = s.strip("` ")
        if s.lower().startswith("json"): s = s[4:].strip()
    try:
        return json.loads(s)
    except Exception:
        try:
            start = s.find("{"); end = s.rfind("}")
            if start != -1 and end != -1 and end > start:
                return json.loads(s[start:end+1])
        except Exception:
            return None
    return None

def llm_turn(user_text: str, state: dict) -> dict:
    """Orquestador con fallback si no hay OpenAI o si da JSON inválido."""
    lite_state = {
        "espacio": state.get("espacio"),
        "modo": state.get("modo"),
        "temperatura": state.get("temperatura"),
        "vatios": state.get("vatios"),
        "instalacion": state.get("instalacion"),
        "presupuesto": state.get("presupuesto"),
    }

    if client is None:
        reply = "Listo, ¿en qué espacio sería la instalación (piscina, oficina, etc.)?"
        return {"reply": reply, "action": "ASK_SPACE", "space": None, "mode": None, "spec_field": None}

    msgs = [
        {"role": "system", "content": SYSTEM_PROMPT_CONDUCTOR},
        {"role": "user", "content": f"Usuario: {user_text}\nEstado:{json.dumps(lite_state, ensure_ascii=False)}"},
    ]

    try:
        r = client.chat.completions.create(
            model=OPENAI_MODEL, temperature=0.2, max_tokens=200, messages=msgs,
        )
        raw = (r.choices[0].message.content or "").strip()
        data = _parse_json_loose(raw) or {}
        if not data: raise ValueError("invalid_json")
        return {
            "reply": data.get("reply") or "Listo.",
            "action": data.get("action") or "NONE",
            "space": data.get("space"),
            "mode": data.get("mode"),
            "spec_field": data.get("spec_field"),
        }
    except Exception as e:
        log_event("llm_error", error=str(e))
        return {
            "reply": "Listo, ¿en qué espacio sería la instalación (piscina, oficina, etc.)?",
            "action": "ASK_SPACE", "space": None, "mode": None, "spec_field": None,
        }

def es_pedido_de_mas(txt: str) -> bool:
    t = _norm(txt)
    TRIGGERS = (
        "mas", "más",            
        "otra", "otras", "otro", "otros",
        "siguiente", "siguientes",
        "adicional", "adicionales",
        "mas opciones", "mas productos", "mas modelos",
        "otra tanda", "otros modelos", "otras opciones"
    )
    return any(k in t for k in TRIGGERS)



# -------------- ESTADO / CHAT ----------
class ChatIn(BaseModel):
    message: str
    lead: Optional[dict] = None
    session_id: Optional[str] = None

def get_state(session_id: str):
    if session_id not in conversation_state:
        conversation_state[session_id] = {
            "espacio": None, "modo": None, "temperatura": None, "vatios": None,
            "instalacion": None, "tipo": None, "presupuesto": None, "mostrados": set(),
        }
    if not isinstance(conversation_state[session_id]["mostrados"], set):
        conversation_state[session_id]["mostrados"] = set(conversation_state[session_id]["mostrados"] or [])
    return conversation_state[session_id]

def _get_mostrados(st: dict, cat: str) -> set:
    st.setdefault("mostrados_por_cat", {})
    st["mostrados_por_cat"].setdefault(cat or "_none_", set())
    return st["mostrados_por_cat"][cat or "_none_"]


LABELS = {
    "lamp_decorativa": "lámparas decorativas",
    "lamp": "lámparas",
    "tubo": "tubos",
    "panel": "paneles",
    "emergencia": "emergencia",
    "industrial": "iluminación industrial",
    "exterior": "exterior",
    "piscina": "luminarias para piscina", 
}



@app.get("/")
def home():
    index_path = os.path.join(STATIC_DIR, "chatbox.html")
    if os.path.exists(index_path): return FileResponse(index_path)
    return {"ok": True, "app": app.title, "version": app.version}

@app.get("/health")
def health():
    return {"ok": True, "env": APP_ENV, "openai": bool(client is not None),
            "products": len(PRODUCTOS), "version": app.version}

@app.post("/reset")
def reset(in_: ChatIn):
    sid = (in_.session_id or "default").strip() or "default"
    conversation_state.pop(sid, None)
    return {"ok": True, "session_id": sid}

@app.get("/session")
def new_session(): return {"session_id": str(uuid.uuid4())}


@app.post("/chat")
async def chat(in_: ChatIn, request: Request):
    if not in_.message or not in_.message.strip():
        raise HTTPException(status_code=400, detail="message is required")
    if len(in_.message) > MAX_MSG_LEN:
        raise HTTPException(status_code=400, detail="message too long")

    ip = _ip(request)
    _rate_limit(ip)

    # --- Session ---
    session_id = (in_.session_id or request.headers.get("x-session-id") or "default").strip() or "default"
    st = get_state(session_id)

    user = in_.message.strip()
    brain = None

    if user.lower() in ["hola", "buenas", "hey", "qué tal"]:
        conversation_state[session_id] = {
            "espacio": None, "modo": None, "temperatura": None, "vatios": None,
            "instalacion": None, "tipo": None, "presupuesto": None, "mostrados": set(),
        }
        return {
            "reply": "👋 Hola, dime qué espacio quieres iluminar (ej: oficina, pasillo, bodega...)",
            "productos": [],
            "session_id": session_id
        }
    budget = parse_presupuesto(user)
    if isinstance(budget, int):
        st["presupuesto"] = budget


        return {
            "reply": "¿Quieres que te sugiera o tienes una luminaria en específico?",
            "productos": [],
            "session_id": session_id
        }

    if brain is None:
        cat = detect_category(user)
        cheapest = want_cheapest(user)

        if cat or cheapest:
            st["modo"] = "sugerir"
            if cat:
                old_cat = st.get("cat")
                if cat != old_cat:
                    st["cat"] = cat
                    st["mostrados"].clear()

            label = LABELS.get(cat, "opciones")

            if cheapest:
                brain = {
                    "reply": f"Te muestro las {label} más económicas 👌",
                    "action": "SEARCH_SPECIFIC",
                    "space": st.get("espacio"),
                    "mode": "sugerir",
                    "spec_field": None,
                }
            else:
                brain = {
                    "reply": f"Listo, te muestro {label} 👌",
                    "action": "SHOW_SUGGESTIONS",
                    "space": st.get("espacio"),
                    "mode": "sugerir",
                    "spec_field": None,
                }

    if brain is None:
        if st.get("espacio") and not st.get("modo"):
            if said_suggest(user):
                st["modo"] = "sugerir"
                brain = {
                    "reply": "Listo, te muestro 5 opciones 👌",
                    "action": "SHOW_SUGGESTIONS",
                    "space": st["espacio"],
                    "mode": "sugerir",
                    "spec_field": None
                }
            elif said_specific(user) or code_regex.search(user) or CODE_SOFT_RE.search(user) or extraer_codigo(user):
                st["modo"] = "especifico"
                brain = {
                    "reply": "Busco esa referencia y te muestro lo que encontré 👌",
                    "action": "SEARCH_SPECIFIC",
                    "space": st["espacio"],
                    "mode": "especifico",
                    "spec_field": None
                }
            else:
                brain = {
                    "reply": "¿Quieres que te sugiera o tienes una luminaria en específico?",
                    "action": "ASK_MODE",
                    "space": st["espacio"],
                    "mode": None,
                    "spec_field": None
                }
        else:
            brain = llm_turn(user, st)

    if brain.get("space") and not st.get("espacio"):
        st["espacio"] = _norm(brain["space"]) or None
    if brain.get("mode") and not st.get("modo"):
        st["modo"] = brain["mode"]

    if not st.get("espacio"):
        t = _norm(user)
        for k in ["piscina","oficina","pasillo","bodega","sala","cocina","baño","banio","terraza","jardin","jardín","fachada","parqueadero","retail","industrial","dormitorio","escalera"]:
            if k in t:
                st["espacio"] = "baño" if k == "banio" else ("jardin" if k == "jardín" else k)
                if brain.get("action") == "ASK_SPACE":
                    brain = {"reply": "¿Quieres que te sugiera o tienes una luminaria en específico?",
                             "action": "ASK_MODE", "space": st["espacio"], "mode": st.get("modo"), "spec_field": None}
                break

    new_cat = detect_category(user)
    if new_cat and new_cat != st.get("cat"):
        st["cat"] = new_cat
        st["mostrados"].clear()
    else:
        new_space = detect_space_from_text(user)

    if brain.get("action") == "COLLECT_SPECS":
        tval = parse_temperatura(user); wval = parse_vatios(user); ival = parse_instalacion(user)
        if tval: st["temperatura"] = tval
        if wval: st["vatios"] = wval
        if ival: st["instalacion"] = ival

    productos: List[dict] = []
    action = brain.get("action", "NONE")

    must_have: Optional[set[str]] = None
    only_one = False
    code_tok: Optional[str] = None

    if cat and cat in CATEGORY_SYNONYMS:
        must_have = {_norm(s) for s in CATEGORY_SYNONYMS[cat]}


    if not cat and st.get("espacio") == "bodega":
        cat = "industrial"
        st["cat"] = cat


    # --- dentro de /chat, justo donde manejas las acciones ---
    if action == "SMALLTALK":
        # No mostrar productos, solo la respuesta ligera
        productos = []
        brain["reply"] = brain.get("reply") or "😊 Claro, pero cuéntame también en qué espacio necesitas la iluminación."

    elif action == "SHOW_SUGGESTIONS":
        cat = st.get("cat")
        label = LABELS.get(cat, "opciones")
        # query semilla: preferir cat/space dinámicamente
        seed = " ".join(sorted(tokens_para_clave(st.get("espacio")) | tokens_para_clave(cat))) or (st.get("espacio") or cat or "")
        productos = buscar_hibrido(seed, st.get("espacio"), cat, k=5)

        if productos:
            _get_mostrados(st, cat).update(p["code"] for p in productos)
            brain["reply"] = f"Listo, te muestro {label} 👌"
        else:
            brain["reply"] = "Por ahora no tengo opciones para ese espacio. ¿Buscas algo específico o te muestro otra categoría? 😉"

    elif action == "SEARCH_SPECIFIC":
        productos = []
        code_tok = extraer_codigo(user)

        # --- 1) Si el usuario dio un código exacto ---
        if code_tok and code_tok in PRODUCTOS:
            p = PRODUCTOS[code_tok]
            brain["reply"] = f"Esto es lo que encontré para {code_tok} 👇"
            return {
                "reply": brain["reply"],
                "productos": [{
                    "code": code_tok,
                    "name": p.get("name"),
                    "price": p.get("price"),
                    "img_url": p.get("img_url"),
                    "url": p.get("url"),
                    "tags": p.get("tags", []),
                    "categories": p.get("categories", []),
                }],
                "session_id": session_id
            }

        # --- 2) Si no coincide directo, buscar aproximado en nombre/categorías ---
        if code_tok:
            for c, p in PRODUCTOS.items():
                cu = c.upper()
                name_u = (p.get("name") or "").upper()
                if code_tok in cu or code_tok in name_u:
                    productos = [{
                        "code": c,
                        "name": p.get("name"),
                        "price": p.get("price"),
                        "img_url": p.get("img_url"),
                        "url": p.get("url"),
                        "tags": p.get("tags", []),
                        "categories": p.get("categories", []),
                    }]
                    brain["reply"] = f"Esto es lo que encontré para {code_tok} 👇"
                    return {"reply": brain["reply"], "productos": productos, "session_id": session_id}

        # --- 3) Fallback: búsqueda híbrida (tokens + embeddings + categoría) ---
        productos = buscar_hibrido(user, st.get("espacio"), st.get("cat"), k=5)

        if productos:
            if code_tok:
                brain["reply"] = f"Esto es lo que encontré para {code_tok} 👇"
            else:
                brain["reply"] = "Esto es lo que encontré 👇"
        else:
            brain["reply"] = "No encontré productos que coincidan con ese código 🤔"


    elif action == "MORE_SUGGESTIONS" or es_pedido_de_mas(user):
        cat = st.get("cat")
        label = LABELS.get(cat, "opciones")

        exclude = _get_mostrados(st, cat)
        productos = sugerencias_aleatorias(n=5, exclude_codes=exclude, categoria=cat)

        if not productos:
            productos = sugerencias_aleatorias(n=5, exclude_codes=exclude, categoria=None)

        if productos:
            exclude.update(p["code"] for p in productos)
            brain["reply"] = f"Te muestro 5 más de {label} 👌"
        else:
            brain["reply"] = f"Ya viste todas las opciones de {label} 😉"


    if code_tok:
        if code_tok in PRODUCTOS:
            p = PRODUCTOS[code_tok]
            return {
                "reply": f"Esto es lo que encontré para {code_tok} 👇",
                "productos": [{
                    "code": code_tok,
                    "name": p.get("name"),
                    "price": p.get("price"),
                    "img_url": p.get("img_url"),
                    "url": p.get("url"),
                    "tags": p.get("tags", []),
                    "categories": p.get("categories", []),
                }],
                "session_id": session_id
            }

        else:
            for c, p in PRODUCTOS.items():
                cu = c.upper()
                name_u = (p.get("name") or "").upper()
                if code_tok in cu or code_tok in name_u:
                    productos = [{
                        "code": c,
                        "name": p.get("name"),
                        "price": p.get("price"),
                        "img_url": p.get("img_url"),
                        "url": p.get("url"),
                        "tags": p.get("tags", []),
                        "categories": p.get("categories", []),
                    }]
                    brain["reply"] = f"Esto es lo que encontré para {code_tok} 👇"
                    return {"reply": brain["reply"], "productos": productos, "session_id": session_id}


        # --- 2) Si no hubo producto por código, pasamos a tokens ---
        if not productos:
            norm_q = _norm(user)
            tokens = [
                _normalize_token(t)
                for t in norm_q.split()
                if len(t) > 2 and t not in {"para", "con", "los", "las", "una", "unos", "unas", "del", "por"}
            ]

            candidatos = []
            for code, p in PRODUCTOS.items():
                full = _norm(
                    (p.get("name") or "") + " " +
                    " ".join(p.get("tags", []) or []) + " " +
                    " ".join(p.get("categories", []) or [])
                )

                # Si el espacio es piscina, exige señales fuertes
                if st.get("espacio") == "piscina" and not (("sumergible" in full) or ("ip68" in full)):
                    continue

                matches = sum(1 for tok in tokens if tok in full)
                if matches >= 1:
                    candidatos.append({
                        "code": code,
                        "name": p.get("name"),
                        "price": p.get("price"),
                        "img_url": p.get("img_url"),
                        "url": p.get("url"),
                        "tags": p.get("tags", []),
                        "categories": p.get("categories", []),
                    })

            if candidatos:
                productos = candidatos[:5]

        # --- 3) Embeddings si los tokens no dieron ---
        if not productos:
            try:
                productos = buscar_por_embeddings(user, k=5)
            except Exception as e:
                log_event("embed_search_error", error=str(e))

        # --- 4) Fallback por categoría detectada / señales del texto ---
        if not productos:
            cat = detect_category(user)
            if cat in CATEGORY_SYNONYMS:
                must_any = {_norm(s) for s in CATEGORY_SYNONYMS[cat]}
            else:
                must_any = set()

            only_one = is_cheapest_question(user)

            # búsqueda específica como último recurso (respeta espacio y cantidad)
            productos = buscar_producto_especifico(
                query_text=user,
                space=st.get("espacio"),
                k=(1 if only_one else 5),
            )

            # si aún vacío y tenemos categoría, intenta un barrido rápido por categoría
            if not productos and must_any:
                productos = buscar_productos(
                    query_text=user,
                    space=st.get("espacio"),
                    k=5,
                    must_any=must_any
                )

            # armar reply en función de lo hallado
            if productos and code_tok:
                brain["reply"] = f"Esto es lo que encontré para {code_tok} 👇"
            elif productos and cat:
                label = LABELS.get(cat, cat)
                brain["reply"] = f"Esto es lo que encontré en {label} 👇"
            elif productos:
                brain["reply"] = "Esto es lo que encontré 👇"


    if action != "SMALLTALK" and not productos and must_have:
        cat = st.get("cat")
        productos = buscar_hibrido(user, st.get("espacio"), cat, k=5)

        for p in productos:
            st["mostrados"].add(p["code"])

        if productos:
            brain["reply"] = "Esto es lo que encontré 👇"
        else:
            brain["reply"] = "No encontré productos que coincidan con eso 🤔"



    if not productos and must_have:
        candidatos = []
        for c, p in PRODUCTOS.items():
            texto = (
                (p.get("name") or "").lower() + " " +
                " ".join((p.get("tags") or [])) + " " +
                " ".join((p.get("categories") or []))
            )
            if any(word in texto for word in must_have):
                candidatos.append({
                    "code": c,
                    "name": p.get("name"),
                    "price": p.get("price"),
                    "img_url": p.get("img_url"),
                    "url": p.get("url"),
                    "tags": p.get("tags", []),
                    "categories": p.get("categories", []),
                })

        if candidatos:
            if want_cheapest(user) or only_one:
                candidatos.sort(key=lambda x: _price_value(x["price"]) or 10**12)
            productos = candidatos[: (1 if only_one else 5)]


        for p in productos:
            st["mostrados"].add(p["code"]) 

        if only_one and productos:
            p0 = productos[0]
            brain["reply"] = f"El más barato es {p0.get('name')} ({p0.get('code')}) por {p0.get('price')} 👇"
        elif want_cheapest(user):
            if cat == "driver":
                brain["reply"] = "Te muestro las fuentes de poder más económicas 👇"
            elif cat == "lamp":
                brain["reply"] = "Te muestro las lámparas más económicas 👇"
            elif cat == "seguridad":
                brain["reply"] = "Te muestro las cámaras de seguridad más económicas 👇"
            else:
                brain["reply"] = "Te muestro las opciones más económicas 👇"
        elif code_tok:
            brain["reply"] = f"Esto es lo que encontré para {code_tok} 👇"
        elif cat:
            brain["reply"] = f"Esto es lo que encontré en {cat} 👇"
        else:
            brain["reply"] = "Te muestro las coincidencias que encontré 👇"


    log_event("chat_turn", session_id=session_id, ip=ip, action=action,
              espacio=st.get("espacio"), modo=st.get("modo"),
              specs={"t": st.get("temperatura"), "w": st.get("vatios"), "i": st.get("instalacion")},
              productos=len(productos))

    return {"reply": brain.get("reply", "Listo."), "productos": productos, "session_id": session_id}

